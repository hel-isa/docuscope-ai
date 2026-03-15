from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image

from app.parsers.docx_parser import parse_docx
from app.parsers.image_parser import parse_image
from app.parsers.pdf_parser import parse_pdf
from app.parsers.txt_parser import parse_txt
from app.parsers.xlsx_parser import parse_xlsx


class _FakePdfPage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _FakePdfReader:
    def __init__(self, _path: str, text: str, metadata: dict | None = None) -> None:
        self.pages = [_FakePdfPage(text)]
        self.metadata = metadata or {}


def test_parse_txt_smoke(tmp_path: Path) -> None:
    f = tmp_path / "sample.txt"
    f.write_text("hello world", encoding="utf-8")

    out = parse_txt(f)

    assert out["text"] == "hello world"
    assert out["page_count"] == 1
    assert out["ocr_used"] is False


def test_parse_docx_smoke(tmp_path: Path) -> None:
    f = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Line 1")
    doc.add_paragraph("Line 2")
    doc.core_properties.author = "Unit Tester"
    doc.save(str(f))

    out = parse_docx(f)

    assert "Line 1" in out["text"]
    assert "Line 2" in out["text"]
    assert out["author_safe"] == "Unit Tester"
    assert out["ocr_used"] is False


def test_parse_xlsx_smoke(tmp_path: Path) -> None:
    f = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["name", "amount"])
    ws.append(["Alice", 100])
    wb.save(str(f))

    out = parse_xlsx(f)

    assert "name | amount" in out["text"]
    assert "Alice | 100" in out["text"]
    assert out["page_count"] == 1
    assert out["tables_detected"] is True


def test_parse_image_smoke_with_mocked_ocr(tmp_path: Path, monkeypatch) -> None:
    f = tmp_path / "sample.png"
    Image.new("RGB", (12, 8), color=(255, 255, 255)).save(f)

    monkeypatch.setattr("app.parsers.image_parser.ocr_image", lambda _path: "ocr text")

    out = parse_image(f)

    assert out["text"] == "ocr text"
    assert out["page_count"] == 1
    assert out["ocr_used"] is True
    assert out["embedded_metadata"]["format"] == "PNG"


def test_parse_pdf_smoke_without_ocr(tmp_path: Path, monkeypatch) -> None:
    f = tmp_path / "sample.pdf"
    f.write_bytes(b"%PDF-1.4\n")

    fake_reader = lambda _path: _FakePdfReader(
        _path,
        text="This is enough embedded PDF text to avoid OCR fallback in parser.",
        metadata={"/Author": "Parser Test"},
    )
    monkeypatch.setattr("app.parsers.pdf_parser.PdfReader", fake_reader)

    out = parse_pdf(f)

    assert out["ocr_needed"] is False
    assert out["ocr_used"] is False
    assert out["author_safe"] == "Parser Test"
    assert "embedded PDF text" in out["text"]


def test_parse_pdf_uses_ocr_when_text_is_short(tmp_path: Path, monkeypatch) -> None:
    f = tmp_path / "short.pdf"
    f.write_bytes(b"%PDF-1.4\n")

    fake_reader = lambda _path: _FakePdfReader(_path, text="tiny", metadata={})
    monkeypatch.setattr("app.parsers.pdf_parser.PdfReader", fake_reader)
    monkeypatch.setattr("app.parsers.pdf_parser.ocr_pdf", lambda _path: "ocr recovered text")

    out = parse_pdf(f)

    assert out["ocr_needed"] is True
    assert out["ocr_used"] is True
    assert out["text"] == "ocr recovered text"
